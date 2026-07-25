"""Generate realistic test data for manual app testing.

Creates a messy folder structure with ~45 files across ~15 folders.
Every file has real content inside. Covers all 8 default rule types,
includes duplicates, protected folders, old files, and edge cases.

Usage:
    python tests/fixtures/generate_test_data.py
    # Then point the app at D:\\test-data\\
"""
import io
import os
import shutil
import struct
import tarfile
import zipfile
from datetime import datetime
from pathlib import Path

# MINIMAL VALID FILE HEADERS (BYTES LITERALS)

# 1X1 WHITE JPEG (SOI AND APP0 AND SOS AND EOI)
JPEG_1X1 = (
    b"\xff\xd8\xff\xe0\x00\x10JFIF\x00\x01\x01\x00\x00\x01\x00\x01\x00\x00"
    b"\xff\xdb\x00C\x00\x08\x06\x06\x07\x06\x05\x08\x07\x07\x07\t\t"
    b"\x08\n\x0c\x14\r\x0c\x0b\x0b\x0c\x19\x12\x13\x0f\x14\x1d\x1a"
    b"\x1f\x1e\x1d\x1a\x1c\x1c $.\x27 ',#\x1c\x1c(7),01444\x1f'9=82<.342"
    b"\xff\xc0\x00\x0b\x08\x00\x01\x00\x01\x01\x01\x11\x00"
    b"\xff\xc4\x00\x1f\x00\x00\x01\x05\x01\x01\x01\x01\x01\x01\x00\x00"
    b"\x00\x00\x00\x00\x00\x00\x01\x02\x03\x04\x05\x06\x07\x08\t\n\x0b"
    b"\xff\xda\x00\x08\x01\x01\x00\x00?\x00\x7f\x80"
    b"\xff\xd9"
)

# 1X1 WHITE PNG (VALID IHDR AND IDAT AND IEND)
def _make_png() -> bytes:
    import zlib
    def _chunk(chunk_type: bytes, data: bytes) -> bytes:
        c = chunk_type + data
        crc = struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
        return struct.pack(">I", len(data)) + c + crc

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    raw = b"\x00\xff\xff\xff"
    compressed = zlib.compress(raw)
    return sig + _chunk(b"IHDR", ihdr_data) + _chunk(b"IDAT", compressed) + _chunk(b"IEND", b"")

PNG_1X1 = _make_png()

# 1X1 WHITE BMP (14 BYTE FILE HEADER AND 40 BYTE DIB HEADER AND PIXEL DATA)
BMP_1X1 = (
    b"BM"
    + struct.pack("<I", 14 + 40 + 12)  # FILE SIZE
    + b"\x00\x00\x00\x00"
    + struct.pack("<I", 14 + 40)
    + struct.pack("<I", 40)
    + struct.pack("<i", 1)
    + struct.pack("<i", 1)
    + struct.pack("<H", 1)
    + struct.pack("<H", 24)
    + b"\x00" * 24
    + b"\xff\xff\xff"
)

# 1X1 WHITE TIFF (LITTLE ENDIAN, BASELINE)
TIFF_1X1 = (
    b"II"  # BYTE ORDER
    + struct.pack("<H", 42)
    + struct.pack("<I", 8)  # IFD OFFSET
    + struct.pack("<H", 12)  # 12 ENTRIES
    + struct.pack("<HHI", 256, 1, 1)  # IMAGE WIDTH
    + struct.pack("<HHI", 257, 1, 1)  # IMAGE LENGTH
    + struct.pack("<HHI", 258, 1, 8)  # BITS PER SAMPLE
    + struct.pack("<HHI", 259, 1, 1)  # COMPRESSION NONE
    + struct.pack("<HHI", 262, 1, 2)  # PHOTOMETRIC INTERPRETATION RGB
    + struct.pack("<HHI", 273, 1, 1)  # STRIP OFFSETS
    + struct.pack("<HHI", 277, 1, 1)  # SAMPLES PER PIXEL
    + struct.pack("<HHI", 278, 1, 1)  # ROWS PER STRIP
    + struct.pack("<HHI", 279, 1, 3)  # STRIP BYTE COUNTS
    + struct.pack("<HHI", 282, 1, 72)  # X RESOLUTION
    + struct.pack("<HHI", 283, 1, 72)  # Y RESOLUTION
    + struct.pack("<HHI", 296, 1, 2)  # RESOLUTION UNIT INCH
    + struct.pack("<I", 0)  # NEXT IFD OFFSET (NONE)
    + b"\xff\xff\xff"  # PIXEL DATA
)

# MINIMAL VALID MP3 FRAME (ID3V2 HEADER AND ONE VALID FRAME)
MP3_1X1 = (
    b"ID3"
    + b"\x03\x00"  # VERSION 2.3
    + b"\x00"  # FLAGS
    + b"\x00\x00\x00\x00"  # SIZE (SYNCSAFE INT, 0)
    + b"\xff\xfb\x90\x00"  # MPEG1 LAYER3 VALID FRAME HEADER
    + b"\x00" * 413  # FRAME DATA PADDING
)

# MINIMAL VALID WAV HEADER
WAV_HEADER = (
    b"RIFF"
    + struct.pack("<I", 36 + 4)  # FILE SIZE MINUS 8
    + b"WAVE"
    + b"fmt "
    + struct.pack("<I", 16)  # CHUNK SIZE
    + struct.pack("<HHIIHH", 1, 1, 44100, 88200, 2, 16)  # PCM, MONO, 44100 HZ
    + b"data"
    + struct.pack("<I", 4)  # DATA SIZE
    + b"\x00\x00\x00\x00"  # ONE SAMPLE OF SILENCE
)

# MINIMAL VALID FLAC HEADER
FLAC_HEADER = (
    b"fLaC"
    + b"\x00"  # METADATA BLOCK STREAMINFO
    + b"\x00\x00\x22"  # BLOCK SIZE 34 BYTES
    + b"\x00\x00"  # MIN BLOCK SIZE
    + b"\x00\x00"  # MAX BLOCK SIZE
    + b"\x00\x00\x00"  # MIN FRAME SIZE
    + b"\x00\x00\x00"  # MAX FRAME SIZE
    + b"\xac\x44"  # SAMPLE RATE 44100 (20 BITS) AND CHANNELS (3 BITS) AND BITS (5 BITS)
    + b"\x00\x00\x00\x00\x00\x00"  # TOTAL SAMPLES AND MD5
)

# MINIMAL VALID MP4 (FTYP BOX)
MP4_HEADER = (
    b"\x00\x00\x00\x14ftypisom"  # FTYP BOX (20 BYTES)
    + b"isom"
    + b"iso2"
    + b"mp41"
)

# MINIMAL VALID MKV (EBML HEADER)
MKV_HEADER = (
    b"\x1a\x45\xdf\xa3"  # EBML MAGIC
    + b"\x01\x00\x00\x00\x00\x00\x00\x0f"  # EBML HEADER ELEMENT
    + b"\x42\x86\x81\x01"  # EBML VERSION 1
    + b"\x42\x87\x81\x04"  # EBML READ VERSION 4
    + b"\x42\x85\x81\x02"  # EBML MAX ID LENGTH 2
    + b"\x42\x83\x81\x08"  # EBML MAX SIZE LENGTH 8
)

# FAKE EXE (MZ HEADER AND PE HEADER)
EXE_HEADER = b"MZ" + b"\x00" * 58 + struct.pack("<I", 64) + b"PE\x00\x00" + b"\x00" * 20

# FAKE DMG HEADER
DMG_HEADER = b"koly" + b"\x00" * 508 + struct.pack("<I", 512)

# FAKE DEB HEADER (AR FORMAT)
DEB_HEADER = b"!<arch>\n" + b"debian-binary   " + b"17000000000\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x30\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x60\x00"


# CONTENT GENERATORS

def _write_text(path: Path, content: str) -> None:
    path.write_text(content, encoding="utf-8")


def _create_documents(folder: Path) -> None:
    folder.mkdir(parents=True, exist_ok=True)

# PDF WITH READABLE TEXT USING PYPDF
    try:
        from pypdf import PdfWriter
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        # ADD ANNOTATIONS WITH TEXT CONTENT
        from pypdf.generic import ArrayObject, DictionaryObject, NameObject, NumberObject, TextStringObject
        annot = DictionaryObject()
        annot[NameObject("/Type")] = NameObject("/Annot")
        annot[NameObject("/Subtype")] = NameObject("/Widget")
        annot[NameObject("/FT")] = NameObject("/Tx")
        annot[NameObject("/V")] = TextStringObject(
            "Invoice #INV-2024-0315\n"
            "Date: March 15, 2024\n"
            "From: Acme Corporation\n"
            "Amount: $1,250.00\n"
            "Description: Quarterly consulting services\n"
            "Payment terms: Net 30"
        )
        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)
        (folder / "invoice_march_2024.pdf").write_bytes(buf.read())
    except ImportError:
        # FALLBACK: WRITE A TEXT FILE WITH .PDF EXTENSION
        _write_text(
            folder / "invoice_march_2024.pdf",
            "Invoice #INV-2024-0315\nDate: March 15, 2024\n"
            "From: Acme Corporation\nAmount: $1,250.00\n"
            "Description: Quarterly consulting services\nPayment terms: Net 30",
        )

    # DOCX WITH REAL PARAGRAPHS
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Project Specification", level=1)
        doc.add_paragraph(
            "This document outlines the technical specification for the "
            "data lifecycle management system. The system must support "
            "offline-first operation with no cloud dependencies."
        )
        doc.add_heading("Requirements", level=2)
        doc.add_paragraph("1. Cross-platform desktop application (Windows, macOS, Linux)")
        doc.add_paragraph("2. SHA-256 duplicate detection")
        doc.add_paragraph("3. Folder classification into 7 categories")
        doc.add_paragraph("4. Rules-based file organization")
        doc.add_paragraph("5. Transaction journal with undo capability")
        doc.save(str(folder / "report_draft_2024.docx"))
    except ImportError:
        _write_text(
            folder / "report_draft_2024.docx",
            "Project Specification\n\nThis document outlines the technical "
            "specification for the data lifecycle management system.\n\n"
            "Requirements:\n1. Cross-platform desktop application\n"
            "2. SHA-256 duplicate detection\n3. Folder classification\n"
            "4. Rules-based organization\n5. Transaction journal with undo",
        )

    # XLSX WITH DATA
    try:
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Budget 2023"
        ws.append(["Category", "Q1", "Q2", "Q3", "Q4", "Total"])
        ws.append(["Infrastructure", 12000, 13500, 14000, 15000, 54500])
        ws.append(["Software Licenses", 8000, 8000, 8500, 8500, 33000])
        ws.append(["Personnel", 45000, 47000, 48000, 50000, 190000])
        ws.append(["Training", 3000, 2500, 4000, 3500, 13000])
        ws.append(["Miscellaneous", 1500, 2000, 1800, 2200, 7500])
        ws.append(["TOTAL", 69500, 73000, 76300, 79200, 298000])
        wb.save(str(folder / "budget_2023.xlsx"))
    except ImportError:
        _write_text(
            folder / "budget_2023.xlsx",
            "Category,Q1,Q2,Q3,Q4,Total\nInfrastructure,12000,13500,14000,15000,54500\n"
            "Software Licenses,8000,8000,8500,8500,33000\nPersonnel,45000,47000,48000,50000,190000\n"
            "Training,3000,2500,4000,3500,13000\nMiscellaneous,1500,2000,1800,2200,7500\n"
            "TOTAL,69500,73000,76300,79200,298000",
        )

    _write_text(
        folder / "meeting_notes.txt",
        "Meeting Notes - March 2024\n"
        "=========================\n\n"
        "Attendees: Alice, Bob, Charlie, Diana\n\n"
        "Agenda:\n"
        "1. Review Stage 2 progress\n"
        "2. Discuss Stage 3 transaction journal design\n"
        "3. Plan AI integration timeline\n"
        "4. Assign action items\n\n"
        "Decisions:\n"
        "- Stage 3 will use a holding area approach for safe undo\n"
        "- Transaction journal will be SQLite-based\n"
        "- AI features will be optional (behind import guards)\n\n"
        "Action Items:\n"
        "- Alice: Design transaction journal schema\n"
        "- Bob: Implement collision detection\n"
        "- Charlie: Build undo system\n"
        "- Diana: Update documentation",
    )

    _write_text(
        folder / "project_specification.md",
        "# OpenCoeus Project Specification\n\n"
        "## Overview\n\n"
        "OpenCoeus is an offline-first Data Lifecycle Management desktop application.\n\n"
        "## Architecture\n\n"
        "```\n"
        "UI (PyQt6) -> Engine -> Scanner/Hasher/Classifier\n"
        "                  |-> Database (SQLite)\n"
        "                  |-> Rules Engine\n"
        "                  |-> Document Extractor\n"
        "```\n\n"
        "## Safety Policy\n\n"
        "- Never delete files automatically\n"
        "- Always show preview before changes\n"
        "- Transaction journal for undo\n"
        "- Protected system folders always excluded\n\n"
        "## Roadmap\n\n"
        "1. Stage 1: Safe audit foundation\n"
        "2. Stage 2: Selective review and organisation\n"
        "3. Stage 3: Execute approved changes\n"
        "4. Stage 4: Spreadsheet workflows\n"
        "5. Stage 5: AI-powered intelligence\n",
    )


def _create_photos(folder: Path) -> None:
    folder.mkdir(parents=True, exist_ok=True)
    (folder / "IMG_20240315_142322.jpg").write_bytes(JPEG_1X1)
    (folder / "IMG_20240315_142322 (1).jpg").write_bytes(JPEG_1X1)      # DUPLICATE
    (folder / "vacation_beach.png").write_bytes(PNG_1X1)
    (folder / "screenshot_2024_01_10.png").write_bytes(PNG_1X1)
    (folder / "family_photo.tiff").write_bytes(TIFF_1X1)


def _create_code(folder: Path) -> None:
    folder.mkdir(parents=True, exist_ok=True)

    _write_text(
        folder / "main.py",
        '"""OpenCoeus main entry point."""\n\n'
        "import sys\n"
        "from pathlib import Path\n\n"
        "from opencoeus.config import ScanSettings\n"
        "from opencoeus.engine import ScanEngine\n\n\n"
        "def main() -> int:\n"
        '    """Run the scan engine."""\n'
        "    if len(sys.argv) < 2:\n"
        '        print("Usage: python main.py <folder>")\n'
        "        return 1\n"
        "    root = Path(sys.argv[1])\n"
        "    settings = ScanSettings(root)\n"
        "    engine = ScanEngine(settings)\n"
        "    result = engine.run()\n"
        f'    print(f"Scanned {{len(result.rows)}} files")\n'
        f'    print(f"Found {{result.duplicate_count}} duplicates")\n'
        "    return 0\n\n\n"
        'if __name__ == "__main__":\n'
        "    raise SystemExit(main())\n",
    )

    _write_text(
        folder / "app.js",
        "// OpenCoeus Desktop Application\n"
        "const { app, BrowserWindow } = require('electron');\n\n"
        "function createWindow() {\n"
        "    const win = new BrowserWindow({\n"
        "        width: 1200,\n"
        "        height: 750,\n"
        "        webPreferences: { nodeIntegration: true }\n"
        "    });\n"
        "    win.loadFile('index.html');\n"
        "}\n\n"
        "app.whenReady().then(createWindow);\n\n"
        "app.on('window-all-closed', () => {\n"
        "    if (process.platform !== 'darwin') app.quit();\n"
        "});\n",
    )

    _write_text(
        folder / "styles.css",
        "/* OpenCoeus Dark Theme */\n"
        ":root {\n"
        "    --bg: #14151e;\n"
        "    --surface: #1a1b2e;\n"
        "    --text: #e2e8f0;\n"
        "    --accent: #38bdf8;\n"
        "    --green: #4ade80;\n"
        "    --red: #f87171;\n"
        "}\n\n"
        "body {\n"
        "    background: var(--bg);\n"
        "    color: var(--text);\n"
        "    font-family: 'Inter', sans-serif;\n"
        "    margin: 0;\n"
        "    padding: 0;\n"
        "}\n\n"
        ".sidebar {\n"
        "    background: #111219;\n"
        "    width: 150px;\n"
        "    height: 100vh;\n"
        "    position: fixed;\n"
        "    left: 0;\n"
        "    top: 0;\n"
        "}\n",
    )

    _write_text(
        folder / "index.html",
        "<!DOCTYPE html>\n"
        '<html lang="en">\n'
        "<head>\n"
        '    <meta charset="UTF-8">\n'
        '    <meta name="viewport" content="width=device-width, initial-scale=1.0">\n'
        "    <title>OpenCoeus</title>\n"
        '    <link rel="stylesheet" href="styles.css">\n'
        "</head>\n"
        "<body>\n"
        '    <div class="sidebar">\n'
        "        <h2>OpenCoeus</h2>\n"
        '        <nav>\n'
        '            <a href="#home">Home</a>\n'
        '            <a href="#scan">Scan</a>\n'
        '            <a href="#results">Results</a>\n'
        "        </nav>\n"
        "    </div>\n"
        '    <div class="content">\n'
        "        <h1>Data Lifecycle Manager</h1>\n"
        "        <p>Offline-first file management and deduplication.</p>\n"
        "    </div>\n"
        "</body>\n"
        "</html>\n",
    )

    _write_text(
        folder / "config.json",
        '{\n'
        '    "app_name": "OpenCoeus",\n'
        '    "version": "0.1.0",\n'
        '    "database": {\n'
        '        "type": "sqlite",\n'
        '        "path": "~/.local/state/opencoeus/data.db"\n'
        "    },\n"
        '    "scanner": {\n'
        '        "max_depth": 5,\n'
        '        "chunk_size": 1048576,\n'
        '        "extract_documents": true\n'
        "    },\n"
        '    "rules": {\n'
        '        "enabled": true,\n'
        '        "priority_order": "ascending"\n'
        "    }\n"
        "}\n",
    )

    _write_text(
        folder / "data.yaml",
        "# OpenCoeus Configuration\n"
        "app:\n"
        "  name: OpenCoeus\n"
        "  version: 0.1.0\n"
        "  description: Offline-first data lifecycle management\n\n"
        "database:\n"
        "  type: sqlite\n"
        "  path: ~/.local/state/opencoeus/data.db\n\n"
        "scanner:\n"
        "  max_depth: 5\n"
        "  chunk_size: 1048576\n"
        "  extract_documents: true\n"
        "  follow_symlinks: false\n\n"
        "rules:\n"
        "  enabled: true\n"
        "  default_rules:\n"
        "    - name: Documents\n"
        "      type: extension\n"
        "      extensions: [pdf, docx, xlsx, txt]\n"
        "      destination: '{folder}/Documents/{filename}'\n",
    )


def _create_music(folder: Path) -> None:
    folder.mkdir(parents=True, exist_ok=True)
    (folder / "song_01.mp3").write_bytes(MP3_1X1)
    (folder / "podcast_episode.wav").write_bytes(WAV_HEADER)
    (folder / "background_music.flac").write_bytes(FLAC_HEADER)


def _create_videos(folder: Path) -> None:
    folder.mkdir(parents=True, exist_ok=True)
    (folder / "tutorial_2024.mp4").write_bytes(MP4_HEADER)
    (folder / "screen_recording.mkv").write_bytes(MKV_HEADER)


def _create_archives(folder: Path) -> None:
    folder.mkdir(parents=True, exist_ok=True)

    # REAL ZIP WITH A FILE INSIDE
    zip_path = folder / "project_backup.zip"
    with zipfile.ZipFile(str(zip_path), "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("backup/readme.txt", "This is a backup of the project files.\nCreated: March 2024.\n")
        zf.writestr("backup/config.json", '{"backup": true, "date": "2024-03-15"}\n')
        zf.writestr("backup/notes.txt", "Remember to verify all files after restore.\n")

    # FAKE RAR HEADER
    (folder / "old_photos.rar").write_bytes(b"Rar!\x1a\x07\x00" + b"\x00" * 500)

    # REAL TAR.GZ WITH A FILE INSIDE
    tar_path = folder / "source_code.tar.gz"
    with tarfile.open(str(tar_path), "w:gz") as tar:
        info = tarfile.TarInfo(name="project/main.py")
        content = b'print("Hello from backup")\n'
        info.size = len(content)
        tar.addfile(info, io.BytesIO(content))

        info2 = tarfile.TarInfo(name="project/README.md")
        content2 = b"# Backup\nThis is a backup of the source code.\n"
        info2.size = len(content2)
        tar.addfile(info2, io.BytesIO(content2))


def _create_installers(folder: Path) -> None:
    folder.mkdir(parents=True, exist_ok=True)
    (folder / "setup_app.exe").write_bytes(EXE_HEADER + b"\x00" * 1000)
    (folder / "installer.dmg").write_bytes(DMG_HEADER + b"\x00" * 500)
    (folder / "package.deb").write_bytes(DEB_HEADER + b"\x00" * 500)


def _create_downloads(folder: Path) -> None:
    folder.mkdir(parents=True, exist_ok=True)

    # ANOTHER PDF
    try:
        from pypdf import PdfWriter
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)
        (folder / "random_document.pdf").write_bytes(buf.read())
    except ImportError:
        (folder / "random_document.pdf").write_bytes(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")

    (folder / "mixed_image.bmp").write_bytes(BMP_1X1)

    _write_text(
        folder / "temporary_notes.csv",
        "Name,Email,Phone,Department\n"
        "Alice Johnson,alice@example.com,555-0101,Engineering\n"
        "Bob Smith,bob@example.com,555-0102,Marketing\n"
        "Charlie Brown,charlie@example.com,555-0103,Sales\n"
        "Diana Prince,diana@example.com,555-0104,Engineering\n"
        "Eve Wilson,eve@example.com,555-0105,HR\n",
    )


def _create_protected_folders(root: Path) -> None:
    # NODE_MODULES (SHOULD BE AUTO EXCLUDED)
    nm = root / "node_modules" / "some_package"
    nm.mkdir(parents=True, exist_ok=True)
    _write_text(
        nm / "index.js",
        'module.exports = {\n'
        '    name: "some-package",\n'
        '    version: "1.0.0",\n'
        '    main: "index.js",\n'
        '    dependencies: {}\n'
        '};\n',
    )
    _write_text(nm / "package.json", '{"name": "some-package", "version": "1.0.0"}\n')

    # .GIT (SHOULD BE AUTO EXCLUDED)
    git_dir = root / ".git"
    git_dir.mkdir(parents=True, exist_ok=True)
    _write_text(
        git_dir / "config",
        "[core]\n"
        "\trepositoryformatversion = 0\n"
        "\tfilemode = true\n"
        "\tbare = false\n"
        "\tlogallrefupdates = true\n"
        "[remote \"origin\"]\n"
        "\turl = https://github.com/user/opencoeus.git\n"
        "\tfetch = +refs/heads/*:refs/remotes/origin/*\n",
    )
    _write_text(git_dir / "HEAD", "ref: refs/heads/main\n")

    # VENV (SHOULD BE AUTO EXCLUDED)
    venv_dir = root / "venv" / "Lib" / "site-packages"
    venv_dir.mkdir(parents=True, exist_ok=True)
    _write_text(venv_dir / "__init__.py", "# Virtual environment packages\n")


def _create_duplicates(root: Path) -> None:
    content = (
        "This is a duplicate test file.\n"
        "It contains some unique text that should be duplicated exactly.\n"
        "Line 3: The quick brown fox jumps over the lazy dog.\n"
        "Line 4: 0123456789 abcdefghijklmnopqrstuvwxyz\n"
        "Line 5: This file will be copied to test duplicate detection.\n"
    )
    _write_text(root / "duplicate_alpha.txt", content)
    _write_text(root / "duplicate_alpha (copy).txt", content)
    _write_text(root / "duplicate_alpha (copy 2).txt", content)


def _create_old_files(root: Path) -> None:
    # CREATE OLD DOCX
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Old Report 2022", level=1)
        doc.add_paragraph(
            "This is an old report from 2022. It should be picked up by the "
            "date-based rule (older_than_days: 365)."
        )
        doc.save(str(root / "old_report_2022.docx"))
    except ImportError:
        _write_text(
            root / "old_report_2022.docx",
            "Old Report 2022\n\nThis is an old report from 2022.",
        )

    # CREATE OLD CSV
    _write_text(
        root / "legacy_data_2021.csv",
        "ID,Date,Amount,Description\n"
        "1,2021-01-15,250.00,Office supplies\n"
        "2,2021-03-22,1200.00,Server maintenance\n"
        "3,2021-06-10,89.99,Software subscription\n"
        "4,2021-09-05,450.00,Travel expenses\n"
        "5,2021-12-20,3200.00,Hardware upgrade\n",
    )

    # SET FILE TIMESTAMPS TO 2022 AND 2021
    old_time_2022 = datetime(2022, 6, 15, 10, 30, 0).timestamp()
    old_time_2021 = datetime(2021, 11, 1, 14, 0, 0).timestamp()
    os.utime(str(root / "old_report_2022.docx"), (old_time_2022, old_time_2022))
    os.utime(str(root / "legacy_data_2021.csv"), (old_time_2021, old_time_2021))


def _create_edge_cases(root: Path) -> None:
    # RANDOM BINARY DATA
    (root / "mixed_file.bin").write_bytes(bytes(range(256)) * 4)

    # FILE WITH SPACES AND SPECIAL CHARACTERS IN NAME
    _write_text(root / "file with spaces.txt", "This file has spaces in its name.\n")

    # FILE WITH UNICODE CHARACTERS
    _write_text(root / "unicode_日本語.txt", "This file contains unicode: 日本語テスト\n")

    # VERY LONG FILENAME (BUT WITHIN OS LIMITS)
    _write_text(root / ("a" * 100 + ".txt"), "File with very long name.\n")

    # EMPTY FILE
    (root / "empty_file.txt").write_bytes(b"")


# MAIN

def create_test_data(root: Path = Path("D:/test-data")) -> None:
    """Create the full test data structure."""
    if root.exists():
        shutil.rmtree(str(root))
    root.mkdir(parents=True)

    print(f"Creating test data at {root}...")

    _create_documents(root / "Documents")
    print("  [+] Documents/ (5 files: PDF, DOCX, XLSX, TXT, MD)")

    _create_photos(root / "Photos")
    print("  [+] Photos/ (5 files: 2x JPG, PNG, PNG, TIFF)")

    _create_code(root / "Code")
    print("  [+] Code/ (6 files: PY, JS, CSS, HTML, JSON, YAML)")

    _create_music(root / "Music")
    print("  [+] Music/ (3 files: MP3, WAV, FLAC)")

    _create_videos(root / "Videos")
    print("  [+] Videos/ (2 files: MP4, MKV)")

    _create_archives(root / "Archives")
    print("  [+] Archives/ (3 files: ZIP, RAR, TAR.GZ)")

    _create_installers(root / "Installers")
    print("  [+] Installers/ (3 files: EXE, DMG, DEB)")

    _create_downloads(root / "Downloads")
    print("  [+] Downloads/ (3 files: PDF, BMP, CSV)")

    _create_protected_folders(root)
    print("  [+] node_modules/, .git/, venv/ (auto-exclude targets)")

    _create_duplicates(root)
    print("  [+] Duplicates (3 identical TXT files)")

    _create_old_files(root)
    print("  [+] Old files (2022 DOCX, 2021 CSV)")

    _create_edge_cases(root)
    print("  [+] Edge cases (binary, spaces, unicode, long name, empty)")

    # COUNT TOTAL FILES
    total = sum(1 for _ in root.rglob("*") if _.is_file())
    print(f"\nDone! {total} files created across {len(list(root.iterdir()))} top-level folders.")
    print(f"Point the app at: {root}")


if __name__ == "__main__":
    create_test_data()
