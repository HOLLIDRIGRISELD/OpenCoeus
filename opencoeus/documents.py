from __future__ import annotations

import logging
import re
from pathlib import Path

logger = logging.getLogger(__name__)

_DOC_TYPE_PATTERNS: list[tuple[str, re.Pattern]] = [
    ("Invoice", re.compile(r"(invoice|inv[-#\s]?\d|amount\s+d(ue|ate)|billing|payment\s+terms)", re.IGNORECASE)),
    ("Meeting-Notes", re.compile(r"(meeting\s*(notes|minutes|log)|attendees?:|agenda|action\s+items)", re.IGNORECASE)),
    ("Specification", re.compile(r"(specification|spec[:.\s]|technical\s*(spec|requirement)|architecture\s*(overview|diagram))", re.IGNORECASE)),
    ("Report", re.compile(r"(report|summary|analysis|findings|conclusion)", re.IGNORECASE)),
    ("Budget", re.compile(r"(budget|financial|revenue|expenses?|forecast|quarterly)", re.IGNORECASE)),
    ("Contract", re.compile(r"(contract|agreement|terms?\s*(and|of)\s*(service|conditions|use)|legal)", re.IGNORECASE)),
    ("Presentation", re.compile(r"(presentation|slides?|deck)", re.IGNORECASE)),
    ("Readme", re.compile(r"(readme|getting\s+started)", re.IGNORECASE)),
    ("Backup", re.compile(r"(backup|archive|restore|snapshot)", re.IGNORECASE)),
    ("Spreadsheet", re.compile(r"(spreadsheet|worksheet|tab|csv|tabular)", re.IGNORECASE)),
    ("Resume", re.compile(r"(resume|cv|curriculum\s+vitae)", re.IGNORECASE)),
    ("Letter", re.compile(r"(dear\s+\w+|sincerely|yours\s+faithfully)", re.IGNORECASE)),
    ("Memo", re.compile(r"(memo|memorandum|to:?|from:?|re:?|subject:?)", re.IGNORECASE)),
    ("Article", re.compile(r"(article|abstract|introduction|methodology|references)", re.IGNORECASE)),
    ("Proposal", re.compile(r"(proposal|statement\s+of\s+work|scope\s+of\s+work|deliverables)", re.IGNORECASE)),
    ("Manual", re.compile(r"(manual|guide|tutorial|user\s+guide|reference)", re.IGNORECASE)),
    ("Certificate", re.compile(r"(certificate|certification|completion|certified)", re.IGNORECASE)),
    ("Form", re.compile(r"(form|application|registration|questionnaire)", re.IGNORECASE)),
    ("Policy", re.compile(r"(policy|procedure|compliance|regulation)", re.IGNORECASE)),
    ("Checklist", re.compile(r"(checklist|check\s*list|todo|tasks)", re.IGNORECASE)),
    ("Thesis", re.compile(r"(thesis|dissertation|doctoral|master['\u2019]s\s+thesis)", re.IGNORECASE)),
    ("Agenda", re.compile(r"(agenda|schedule|timeline|itinerary)", re.IGNORECASE)),
    ("Newsletter", re.compile(r"(newsletter|news\s*letter|issue\s+\d)", re.IGNORECASE)),
    ("Order", re.compile(r"(order|purchase\s+order|po\s*[-\s]?\d+)", re.IGNORECASE)),
    ("Quote", re.compile(r"(quote|quotation|estimate|price\s+list)", re.IGNORECASE)),
    ("Timesheet", re.compile(r"(timesheet|time\s*sheet|time\s+tracking|hours)", re.IGNORECASE)),
]

# LINES MATCHING THESE PATTERNS ARE LIKELY HEADERS, FOOTERS, OR PAGE NUMBERS
_NOISE_PATTERNS = re.compile(
    r"^(page\s+\d+|第.+页|header|footer|confidential|draft|do not distribute)$",
    re.IGNORECASE,
)

# LINES CONTAINING THESE PHRASES ARE LIKELY TITLE CANDIDATES
_TITLE_INDICATORS = re.compile(
    r"(title|subject|report|invoice|specification|summary|meeting|notes|agenda|minutes)",
    re.IGNORECASE,
)

# LINES STARTING WITH THESE ARE LIKELY HEADERS
_HEADER_PREFIXES = re.compile(
    r"^(#+\s+|[-=]+\s*$|abstract|introduction|overview|background)",
    re.IGNORECASE,
)


def detect_document_type(text: str, metadata: dict | None = None, title_candidate: str = "") -> str:
    """Analyze text and metadata to determine document type.
    Returns a type name like "Invoice", "Meeting-Notes", "Specification", or fallback "Document".
    Checks title candidate first (highest weight), then text body."""
    combined = text
    if title_candidate:
        combined = title_candidate + "\n" + text
    if metadata:
        meta_text = " ".join(str(v) for v in metadata.values() if v)
        combined = combined + "\n" + meta_text
    for doc_type_name, doc_type_pattern in _DOC_TYPE_PATTERNS:
        if doc_type_pattern.search(combined):
            return doc_type_name
    return "Document"


def _extract_pdf_text_annotations(pdf_page) -> str:
    """Fallback: extract text from PDF annotations and form fields."""
    texts: list[str] = []
    try:
        for annot in pdf_page.get("/Annots", []):
            annot_obj = annot.get_object() if hasattr(annot, "get_object") else annot
            if isinstance(annot_obj, dict):
                for key in ("/V", "/Contents", "/T"):
                    value = annot_obj.get(key)
                    if isinstance(value, str) and value.strip():
                        texts.append(value.strip())
    except Exception:
        pass
    return "\n".join(texts)


def extract_text(document_path: Path, maximum_pages: int = 2) -> str:
    logger.debug("Extracting text from %s", document_path)
    try:
        if document_path.suffix.lower() == ".pdf":
            from pypdf import PdfReader
            pdf_reader = PdfReader(str(document_path))
            text_parts: list[str] = []
            for pdf_page in pdf_reader.pages[:maximum_pages] if maximum_pages < 999 else pdf_reader.pages:
                page_text = pdf_page.extract_text(extraction_mode="layout") or ""
                if not page_text.strip():
                    page_text = pdf_page.extract_text() or ""
                if not page_text.strip():
                    page_text = _extract_pdf_text_annotations(pdf_page)
                text_parts.append(page_text)
            return "\n".join(text_parts)
        if document_path.suffix.lower() == ".docx":
            from docx import Document
            word_document = Document(str(document_path))
            paragraphs = word_document.paragraphs[:80] if maximum_pages < 999 else word_document.paragraphs
            return "\n".join(p.text for p in paragraphs)
        if document_path.suffix.lower() in {".txt", ".md"}:
            lines: list[str] = []
            with open(document_path, "r", encoding="utf-8", errors="replace") as text_file:
                if maximum_pages < 999:
                    for line_number, line in enumerate(text_file):
                        if line_number >= 15:
                            break
                        lines.append(line)
                else:
                    lines = text_file.readlines()
            return "".join(lines)
        if document_path.suffix.lower() in {".xlsx", ".xls"}:
            return _extract_xlsx_text(document_path)
        if document_path.suffix.lower() == ".csv":
            return _extract_csv_text(document_path)
    except Exception as exc:
        logger.debug("Text extraction failed for %s: %s", document_path, exc)
        return ""
    return ""


def _extract_xlsx_text(file_path: Path) -> str:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(file_path, read_only=True, data_only=True)
        texts = []
        for sheet_name in wb.sheetnames[:3]:
            ws = wb[sheet_name]
            texts.append(f"Sheet: {sheet_name}")
            for row_idx, row in enumerate(ws.iter_rows(values_only=True)):
                if row_idx >= 20:
                    break
                row_text = " | ".join(str(c) for c in row if c is not None)
                if row_text.strip():
                    texts.append(row_text)
        wb.close()
        return "\n".join(texts)
    except Exception:
        return ""


def _extract_csv_text(file_path: Path) -> str:
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read(65536)
    except Exception:
        return ""


def extract_metadata(document_path: Path) -> dict:
    """Extract metadata from documents: PDF/DOCX author, title, creation date."""
    metadata: dict = {}
    try:
        if document_path.suffix.lower() == ".pdf":
            from pypdf import PdfReader
            pdf_reader = PdfReader(str(document_path))
            info = pdf_reader.metadata
            if info:
                metadata["author"] = info.author or ""
                metadata["title"] = info.title or ""
                metadata["subject"] = info.subject or ""
                metadata["creator"] = info.creator or ""
                if info.creation_date:
                    metadata["created_date"] = info.creation_date.strftime("%Y-%m-%d")
                    metadata["created_month"] = info.creation_date.strftime("%m")
        if document_path.suffix.lower() == ".docx":
            from docx import Document
            word_document = Document(str(document_path))
            try:
                props = word_document.core_properties
                if props.author:
                    metadata["author"] = props.author
                if props.title:
                    metadata["title"] = props.title
                if props.subject:
                    metadata["subject"] = props.subject
                if props.created:
                    metadata["created_date"] = props.created.strftime("%Y-%m-%d")
                    metadata["created_month"] = props.created.strftime("%m")
            except Exception:
                pass
    except Exception as exc:
        logger.debug("Metadata extraction failed for %s: %s", document_path, exc)
    return metadata


def _score_title_candidate(line: str) -> float:
    """Score a line as a title candidate; higher is better."""
    score = 0.0
    words = line.split()
    word_count = len(words)
    # PREFER 2 TO 8 WORDS
    if 2 <= word_count <= 8:
        score += 2.0
    elif 1 <= word_count <= 12:
        score += 1.0
    # PREFER MIXED CASE OVER ALL CAPS OR ALL LOWERCASE
    if any(c.isupper() for c in line) and any(c.islower() for c in line):
        score += 1.5
    elif line.isupper():
        score -= 0.5
    # PENALIZE LINES THAT LOOK LIKE NOISE
    if _NOISE_PATTERNS.match(line.strip()):
        score -= 3.0
    # PREFER LINES WITH LETTERS OVER PURE NUMBERS
    alpha_ratio = sum(c.isalpha() for c in line) / max(len(line), 1)
    if alpha_ratio > 0.5:
        score += 0.5
    # BOOST LINES CONTAINING TITLE INDICATORS LIKE INVOICE REPORT OR SUMMARY
    if _TITLE_INDICATORS.search(line):
        score += 2.0
    # PENALIZE LINES THAT LOOK LIKE MARKDOWN HEADERS OR SECTION TITLES
    if _HEADER_PREFIXES.match(line.strip()):
        score -= 1.0
    # PREFER SHORTER LINES WITHIN THE IDEAL RANGE
    if 20 <= len(line.strip()) <= 80:
        score += 0.5
    return score


def suggest_title(document_text: str, fallback_title: str) -> str:
    """Select a short, readable line from document text as a title suggestion."""
    text_lines = [re.sub(r"\s+", " ", line).strip() for line in document_text.splitlines()]
    candidate_titles = [
        line for line in text_lines
        if 5 <= len(line) <= 120 and any(character.isalpha() for character in line)
    ]
    if not candidate_titles:
        safe_filename_title = re.sub(r'[<>:"/\\|?*]', "-", fallback_title).strip(" .-")
        return safe_filename_title[:120] or "Untitled document"
    # SCORE AND SELECT THE BEST CANDIDATE
    scored_candidates = [(_score_title_candidate(line), line) for line in candidate_titles]
    scored_candidates.sort(key=lambda item: item[0], reverse=True)
    selected_title = scored_candidates[0][1]
    safe_filename_title = re.sub(r'[<>:"/\\|?*]', "-", selected_title).strip(" .-")
    return safe_filename_title[:120] or "Untitled document"
